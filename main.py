from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import datetime
import json
from docx.enum.text import WD_ALIGN_PARAGRAPH

data = json.load(open('data.json'))
# print (data)
document = Document()
currentDT = datetime.datetime.now()
print ("Enter The Position to Apply:")
companyposition = input()
print ("Enter The Company Name:")
companyname = input()

p1 = document.add_heading(data['name'], 0)
p1.add_run('\n')
info_address = p1.add_run(data['contact'])
info_address.font.size = Pt(10)
p1.add_run('\n')
info_github = p1.add_run(data['Github'])
info_github.font.size = Pt(10)

p2 = document.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p2.add_run(currentDT.strftime('%d, %b %Y'))

p3 = document.add_paragraph(data['To_Whom'])
p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

apply = data['Apply']
p4 = document.add_paragraph(apply + companyposition)
p4 = p4.add_run(" at " + companyname + " where I can develop my technical skill and strive to grow my career within the company.")
p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

introduce = data['Introduce']
p5 = document.add_paragraph(introduce)
p5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

my_project = data['Experience_that_fits']
p6 = document.add_paragraph(my_project)
p6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

my_business_background = data['My_characters']
p7 = document.add_paragraph(my_business_background)
p7.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

my_attribute_and_contact1 = data['Thank_you']
my_attribute_and_contact2 = data['Please_Contact_Me']
p8 = document.add_paragraph(my_attribute_and_contact1)
p8.add_run(companyname + ".")
p8.add_run(my_attribute_and_contact2)
p8.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

p9 = document.add_paragraph('Sincerely,')

p10 = document.add_paragraph('Ryan Nguyen')



document.save('demo.docx')

